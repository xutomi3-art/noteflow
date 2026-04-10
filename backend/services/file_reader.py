"""Quick file text extraction for Just Chat (bypasses MinerU/RAGFlow)."""

import logging
import os

logger = logging.getLogger(__name__)


async def extract_text(file_path: str, filename: str) -> str:
    """Extract text from a file. Returns empty string on failure."""
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext in (".txt", ".md", ".csv"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()[:50000]

        elif ext == ".docx":
            from docx import Document
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        tables_text.append(" | ".join(cells))
            return "\n".join(paragraphs + tables_text)[:50000]

        elif ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            lines = []
            for sheet in wb.sheetnames[:5]:  # max 5 sheets
                ws = wb[sheet]
                lines.append(f"--- Sheet: {sheet} ---")
                for row in ws.iter_rows(max_row=200, values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    if any(c.strip() for c in cells):
                        lines.append(" | ".join(cells))
            wb.close()
            return "\n".join(lines)[:50000]

        elif ext == ".pdf":
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text = ""
                for page in doc[:30]:  # max 30 pages
                    text += page.get_text() + "\n"
                doc.close()
                return text[:50000]
            except ImportError:
                logger.warning("PyMuPDF not installed, skipping PDF extraction")
                return ""

        else:
            return ""

    except Exception as e:
        logger.warning("File text extraction failed for %s: %s", filename, e)
        return ""
